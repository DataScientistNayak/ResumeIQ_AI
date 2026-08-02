"""
DOCX Parser Utility
"""

from __future__ import annotations

from docx import Document


def extract_docx_text(docx_file) -> str:
    """
    Extract text from a DOCX resume.
    Includes paragraphs and table content.
    """

    document = Document(docx_file)

    text = []

    # Paragraphs
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text.append(paragraph.text.strip())

    # Tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    text.append(cell_text)

    return "\n".join(text)